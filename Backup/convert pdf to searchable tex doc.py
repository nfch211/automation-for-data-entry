import os
import tkinter as tk
from tkinter import filedialog, messagebox
import pytesseract
from PIL import Image
from docx import Document
from pdf2image import convert_from_path

# Set up the Tkinter window
window = tk.Tk()
window.title("Image-to-Word Converter")

def select_input_folder():
    input_folder = filedialog.askdirectory(title="Select Input Folder")
    if input_folder:
        entry_input_folder.delete(0, tk.END)
        entry_input_folder.insert(tk.END, input_folder)

def select_output_folder():
    output_folder = filedialog.askdirectory(title="Select Output Folder")
    if output_folder:
        entry_output_folder.delete(0, tk.END)
        entry_output_folder.insert(tk.END, output_folder)

def convert_image_pdf_to_word():
    pytesseract.pytesseract.tesseract_cmd = r"C:\Users\hofong\Desktop\python\Tesseract\tesseract.exe"
    input_folder = entry_input_folder.get()
    output_folder = entry_output_folder.get()

    if not input_folder or not output_folder:
        messagebox.showerror("Error", "Please select input and output folders.")
        return

    # Create the output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)

    # Set the path to the Poppler binary directory
    poppler_path = r"C:\Users\hofong\Desktop\python\poppler-23.07.0\Library\bin"
    os.environ["PATH"] += os.pathsep + poppler_path

    # Loop through files in the input folder
    for root, dirs, files in os.walk(input_folder):
        for file in files:
            file_path = os.path.join(root, file)
            file_name, file_ext = os.path.splitext(file)

            # Only process image files
            if file_ext.lower() in [".jpg", ".jpeg", ".png", ".bmp", ".gif"]:
                # Open the image file using PIL
                image = Image.open(file_path)

                # Perform OCR to extract the text
                ocr_text = pytesseract.image_to_string(image, lang="chi_sim")

                # Create a new Word document
                doc = Document()

                # Add the OCR text to the document
                doc.add_paragraph(ocr_text)

                # Save the Word document
                output_file_path = os.path.join(output_folder, file_name + ".docx")
                doc.save(output_file_path)

            elif file_ext.lower() == ".pdf":
                # Convert PDF pages to images
                try:
                    pages = convert_from_path(file_path)
                except Exception as e:
                    messagebox.showwarning("PDF Conversion Error", str(e))
                    continue

                # Create a new Word document
                doc = Document()

                # Loop through the converted images and perform OCR
                for page in pages:
                    # Perform OCR to extract the text from the image
                    ocr_text = pytesseract.image_to_string(page, lang="chi_sim")

                    # Add the OCR text to the document
                    doc.add_paragraph(ocr_text)

                # Save the Word document
                output_file_path = os.path.join(output_folder, file_name + ".docx")
                doc.save(output_file_path)

    messagebox.showinfo("Conversion Complete", "Images and PDFs converted to Word successfully.")

# Create and place the input folder selection widgets
label_input_folder = tk.Label(window, text="Input Folder:")
label_input_folder.pack()

entry_input_folder = tk.Entry(window, width=50)
entry_input_folder.pack()

button_select_input_folder = tk.Button(window, text="Select", command=select_input_folder)
button_select_input_folder.pack()

# Create and place the output folder selection widgets
label_output_folder = tk.Label(window, text="Output Folder:")
label_output_folder.pack()

entry_output_folder = tk.Entry(window, width=50)
entry_output_folder.pack()

button_select_output_folder = tk.Button(window, text="Select", command=select_output_folder)
button_select_output_folder.pack()

# Create the conversion button
button_convert = tk.Button(window, text="Convert", command=convert_image_pdf_to_word)
button_convert.pack()

# Start the Tkinter event loop
window.mainloop()
